import os
import hashlib
from typing import List, Tuple
from docx import Document
from backend.config import settings

class DocumentProcessor:
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a Word document."""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
        except Exception as e:
            raise Exception(f"Error extracting text from document: {str(e)}")
    
    def chunk_text(self, text: str) -> List[Tuple[str, int, int]]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        
        chunks = []
        words = text.split()
        
        if len(words) <= self.chunk_size:
            return [(text, 0, len(text))]
        
        start_idx = 0
        char_start = 0
        
        while start_idx < len(words):
            end_idx = min(start_idx + self.chunk_size, len(words))
            chunk_words = words[start_idx:end_idx]
            chunk_text = " ".join(chunk_words)
            
            # Calculate character positions
            char_end = char_start + len(chunk_text)
            
            chunks.append((chunk_text, char_start, char_end))
            
            # Move start position with overlap
            if end_idx >= len(words):
                break
            
            start_idx = end_idx - self.chunk_overlap
            char_start = char_end - len(" ".join(words[end_idx - self.chunk_overlap:end_idx]))
        
        return chunks
    
    def calculate_content_hash(self, text: str) -> str:
        """Calculate SHA-256 hash of content for deduplication."""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()
    
    def process_document(self, file_path: str) -> Tuple[str, List[Tuple[str, int, int]], str]:
        """Process a document and return text, chunks, and content hash."""
        text = self.extract_text_from_docx(file_path)
        chunks = self.chunk_text(text)
        content_hash = self.calculate_content_hash(text)
        
        return text, chunks, content_hash